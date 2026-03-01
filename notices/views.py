import re
import requests
import openpyxl
from collections import OrderedDict
from datetime import date, timedelta, datetime
from io import BytesIO

from django.conf import settings
from django.db.models import Q
from django.http import JsonResponse, HttpResponse
from django.shortcuts import render, get_object_or_404, redirect
from django.views.decorators.http import require_POST, require_GET
from django.contrib import messages

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from openpyxl.styles import Font, PatternFill
from openpyxl.utils import get_column_letter

from .models import Notice, NoticeItem
from .forms import NoticeForm, NoticeItemFormSet, SearchForm


# ──────────────────────────────────────────────
# Подсветка поиска
# ──────────────────────────────────────────────
def highlight(text, query):
    if not query or not text:
        return text
    pattern = re.compile(re.escape(query), re.IGNORECASE)
    return pattern.sub(lambda m: f'<mark>{m.group(0)}</mark>', str(text))


# ──────────────────────────────────────────────
# Список
# ──────────────────────────────────────────────
def notice_list(request):
    form  = SearchForm(request.GET or None)
    qs    = Notice.objects.prefetch_related('items').all()
    query = ''

    if form.is_valid():
        query      = form.cleaned_data.get('q', '')
        date_from  = form.cleaned_data.get('date_from')
        date_to    = form.cleaned_data.get('date_to')
        date_field = form.cleaned_data.get('date_field') or 'issue_date'

        if query:
            qs = qs.filter(
                Q(newspaper__icontains=query) |
                Q(items__address__icontains=query) |
                Q(items__cadastral_number__icontains=query) |
                Q(items__customer__icontains=query) |
                Q(items__contract__icontains=query)
            ).distinct()

        if date_from:
            qs = qs.filter(**{f'{date_field}__gte': date_from})
        if date_to:
            qs = qs.filter(**{f'{date_field}__lte': date_to})

    notices = []
    for n in qs:
        items_hl = []
        for it in n.items.all():
            items_hl.append({
                'address':          highlight(it.address, query),
                'cadastral_number': highlight(it.cadastral_number, query),
                'customer':         highlight(it.customer, query),
                'contract':         highlight(it.contract, query),
            })
        notices.append({
            'obj':             n,
            'newspaper':       highlight(n.newspaper, query),
            'items':           items_hl,
            'approval_status': n.approval_status,
        })

    return render(request, 'notices/list.html', {
        'notices':     notices,
        'search_form': form,
        'query':       query,
    })


# ──────────────────────────────────────────────
# Создание
# ──────────────────────────────────────────────
def notice_create(request):
    if request.method == 'POST':
        form    = NoticeForm(request.POST)
        formset = NoticeItemFormSet(request.POST)
        if form.is_valid() and formset.is_valid():
            notice           = form.save()
            formset.instance = notice
            formset.save()
            messages.success(request, 'Извещение добавлено.')
            return redirect('notice_list')
        errors = {**form.errors}
        for i, fe in enumerate(formset.errors):
            if fe:
                errors[f'Строка {i+1}'] = list(fe.values())
        return JsonResponse({'success': False, 'errors': errors}, status=400)
    return redirect('notice_list')


# ──────────────────────────────────────────────
# Получение данных для редактирования (AJAX)
# ──────────────────────────────────────────────
@require_GET
def notice_get(request, pk):
    notice = get_object_or_404(Notice, pk=pk)
    items  = list(notice.items.values(
        'id', 'address', 'fias_id', 'region', 'city',
        'street', 'house', 'cadastral_number', 'customer', 'contract', 'order'
    ))
    return JsonResponse({
        'id':            notice.pk,
        'newspaper':     notice.newspaper,
        'issue_date':    notice.issue_date.isoformat()    if notice.issue_date    else '',
        'approval_date': notice.approval_date.isoformat() if notice.approval_date else '',
        'items':         items,
    })


# ──────────────────────────────────────────────
# Редактирование
# ──────────────────────────────────────────────
def notice_edit(request, pk):
    notice = get_object_or_404(Notice, pk=pk)

    if request.method == 'POST':
        form    = NoticeForm(request.POST, instance=notice)
        formset = NoticeItemFormSet(request.POST, instance=notice)

        if form.is_valid() and formset.is_valid():
            form.save()
            notice.items.all().delete()
            instances = formset.save(commit=False)
            for obj in instances:
                obj.notice = notice
                obj.save()
            for obj in formset.deleted_objects:
                if obj.pk:
                    obj.delete()
            messages.success(request, 'Извещение обновлено.')
            return redirect('notice_list')

        errors = {**form.errors}
        for i, fe in enumerate(formset.errors):
            if fe:
                errors[f'Строка {i + 1}'] = list(fe.values())
        return JsonResponse({'success': False, 'errors': errors}, status=400)

    return redirect('notice_list')


# ──────────────────────────────────────────────
# Удаление
# ──────────────────────────────────────────────
@require_POST
def notice_delete(request, pk):
    get_object_or_404(Notice, pk=pk).delete()
    messages.success(request, 'Извещение удалено.')
    return redirect('notice_list')


# ──────────────────────────────────────────────
# ФИАС (DaData)
# ──────────────────────────────────────────────
def fias_suggest(request):
    q = request.GET.get('q', '')
    if not q:
        return JsonResponse({'suggestions': []})
    token = getattr(settings, 'DADATA_TOKEN', '')
    if not token or token == 'ВАШ_ТОКЕН_DADATA':
        return JsonResponse({'suggestions': [
            {'value': q + ' (укажите DADATA_TOKEN)', 'fias_id': '',
             'region': '', 'city': '', 'street': '', 'house': ''}
        ]})
    try:
        resp = requests.post(
            settings.FIAS_API_URL,
            json={'query': q, 'count': 10},
            headers={
                'Authorization': f'Token {token}',
                'Content-Type': 'application/json',
            },
            timeout=3
        )
        resp.raise_for_status()
        out = []
        for s in resp.json().get('suggestions', []):
            d = s.get('data', {})
            out.append({
                'value':   s.get('value', ''),
                'fias_id': d.get('fias_id', '') or d.get('house_fias_id', ''),
                'region':  d.get('region_with_type', ''),
                'city':    d.get('city_with_type', '') or d.get('settlement_with_type', ''),
                'street':  d.get('street_with_type', ''),
                'house':   d.get('house', ''),
            })
        return JsonResponse({'suggestions': out})
    except Exception as e:
        return JsonResponse({'suggestions': [], 'error': str(e)})


# ──────────────────────────────────────────────
# Экспорт Word
# ──────────────────────────────────────────────
_MONTHS_RU = [
    '', 'января', 'февраля', 'марта', 'апреля', 'мая', 'июня',
    'июля', 'августа', 'сентября', 'октября', 'ноября', 'декабря',
]


def notice_export_word(request, pk):
    notice = get_object_or_404(Notice, pk=pk)
    items  = list(notice.items.all())

    doc = Document()

    # ── Страница ──────────────────────────────
    sec               = doc.sections[0]
    sec.page_width    = Mm(210)
    sec.page_height   = Mm(297)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(1.5)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

    sec.header.is_linked_to_previous = False
    sec.footer.is_linked_to_previous = False
    for hf in (sec.header, sec.footer):
        for p in hf.paragraphs:
            p.clear()

    # ── Базовый стиль ─────────────────────────
    ns = doc.styles['Normal']
    ns.font.name      = 'Times New Roman'
    ns.font.size      = Pt(12)
    ns.font.bold      = False
    ns.font.color.rgb = RGBColor(0, 0, 0)

    BLACK = RGBColor(0, 0, 0)

    def run(para, text, bold=False, underline=False,
            size=12, name='Times New Roman'):
        r = para.add_run(text)
        r.font.name      = name
        r.font.size      = Pt(size)
        r.bold           = bold
        r.italic         = False
        r.underline      = underline
        r.font.color.rgb = BLACK
        try:
            rPr    = r._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:cs'),       name)
            rFonts.set(qn('w:eastAsia'), name)
        except Exception:
            pass
        return r

    def para(align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             first_indent=None, left_indent=None,
             space_before=0, space_after=0, line_spacing=14):
        p   = doc.add_paragraph()
        fmt = p.paragraph_format
        fmt.alignment    = align
        fmt.space_before = Pt(space_before)
        fmt.space_after  = Pt(space_after)
        fmt.line_spacing = Pt(line_spacing)
        if first_indent is not None:
            fmt.first_line_indent = Cm(first_indent)
        if left_indent is not None:
            fmt.left_indent = Cm(left_indent)
        return p

    def fmt_date_long(d):
        return f'{d.day} {_MONTHS_RU[d.month]} {d.year}'

    def fmt_date_short(d):
        return d.strftime('%d.%m.%Y')

    PLACEHOLDER      = '__.__.____'
    PLACEHOLDER_LONG = '__ ________ ____'

    # ── Заголовок ─────────────────────────────
    p = para(align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    run(p,
        'ИЗВЕЩЕНИЕ О ПРОВЕДЕНИИ СОБРАНИЯ О СОГЛАСОВАНИИ\n'
        'МЕСТОПОЛОЖЕНИЯ ГРАНИЦ ЗЕМЕЛЬНЫХ УЧАСТКОВ')

    # ── Вводный абзац ─────────────────────────
    p = para(first_indent=1.25, space_before=6)
    run(p, 'Филиалом Публично-правовой компании «')
    run(p, 'Роскадастр', underline=True)
    run(p,
        '» по Луганской Народной Республике, юридический адрес: '
        '291011, РФ, ЛНР, г. Луганск, ул. Шелкового, д. 1д, e-mail: ')
    run(p, 'filial1@81.kadastr.ru', underline=True)
    run(p,
        ', ОКПО 74750952, ОГРН 1227700700633, ИНН/КПП 7708410783/940343001, '
        'контактный телефон (0-22) 50 12 90, в лице заместителя директора '
        'филиала – главного технолога публично-правовой компании «')
    run(p, 'Роскадастр', underline=True)
    run(p, '» по Луганской Народной Республике ')
    run(p, 'Прядкина Гавриила Викторовича', underline=True)
    run(p,
        ' (доверенность от 09.01.2025) выполняются кадастровые работы '
        'в отношении земельных участков, расположенных по адресу '
        '(местонахождение):')

    # ── Список объектов ───────────────────────
    for i, it in enumerate(items, 1):
        p = para(left_indent=1.25, first_indent=0,
                 space_before=2, space_after=2)
        parts = []
        if it.address:
            parts.append(it.address)
        if it.cadastral_number:
            parts.append(f'кадастровый номер {it.cadastral_number}')
        if it.customer:
            parts.append(f'заказчик: {it.customer}')
        line = ', '.join(parts) + (';' if parts else '')
        run(p, f'{i}.\t{line}')

    # ── Дата собрания ─────────────────────────
    p = para(first_indent=1.25, space_before=6)
    run(p,
        'Собрание по поводу согласования местоположения границ состоится '
        'по адресу: РФ, ЛНР, г. Луганск, ул. Оборонная, д. 101Б ')
    approval_str = (fmt_date_short(notice.approval_date)
                    if notice.approval_date else PLACEHOLDER)
    run(p, approval_str)
    run(p, ' в 10 часов 00 минут.')

    # ── Ознакомление ──────────────────────────
    p = para(first_indent=1.25, space_before=6)
    run(p,
        'С проектами межевых планов можно ознакомиться по адресу: '
        'г. Луганск, ул. Оборонная, д. 101Б.')

    # ── Сроки возражений ──────────────────────
    p = para(first_indent=1.25, space_before=6)
    run(p,
        'Требования о проведении согласования местоположения границ земельных '
        'участков на местности и обоснованные возражения о местоположении '
        'границ после ознакомления с проектами межевых планов принимаются с ')
    issue_str = (fmt_date_long(notice.issue_date)
                 if notice.issue_date else PLACEHOLDER_LONG)
    run(p, issue_str)
    run(p, ' по ')
    approval_long = (fmt_date_long(notice.approval_date)
                     if notice.approval_date else PLACEHOLDER_LONG)
    run(p, approval_long)
    run(p, ' по адресу: г. Луганск, ул. Оборонная, д. 101Б.')

    # ── Документы ─────────────────────────────
    p = para(first_indent=1.25, space_before=6)
    run(p,
        'При проведении согласования местоположения границ при себе иметь '
        'документ, удостоверяющий личность, а также документы о правах на '
        'земельный участок (часть 12 статьи 39, часть 2 статьи 40 Федерального '
        'закона от 24 июля 2007 г. № 221-ФЗ «О кадастровой деятельности»).')

    # ── Сохранение ────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    fname = f'izveshenie_{notice.pk}.docx'
    resp  = HttpResponse(
        buf.read(),
        content_type=(
            'application/vnd.openxmlformats-officedocument'
            '.wordprocessingml.document'
        )
    )
    resp['Content-Disposition'] = (
        f'attachment; filename="{fname}"; '
        f"filename*=UTF-8''{fname}"
    )
    return resp


# ──────────────────────────────────────────────
# Вспомогательная функция парсинга дат
# ──────────────────────────────────────────────
def parse_date(value):
    """Пробует распарсить дату из разных форматов."""
    if not value:
        return None
    if isinstance(value, datetime):
        return value.date()
    if hasattr(value, 'date'):
        return value
    for fmt in ('%d.%m.%Y', '%Y-%m-%d', '%d/%m/%Y'):
        try:
            return datetime.strptime(str(value).strip(), fmt).date()
        except ValueError:
            continue
    return None


# ──────────────────────────────────────────────
# Импорт из Excel
# ──────────────────────────────────────────────
@require_POST
def notice_import(request):
    excel_file = request.FILES.get('excel_file')
    if not excel_file:
        return JsonResponse({'success': False, 'error': 'Файл не выбран'}, status=400)

    if not excel_file.name.endswith(('.xlsx', '.xls')):
        return JsonResponse(
            {'success': False, 'error': 'Поддерживаются только .xlsx и .xls'},
            status=400
        )

    try:
        wb = openpyxl.load_workbook(excel_file, data_only=True)
        ws = wb.active
    except Exception as e:
        return JsonResponse(
            {'success': False, 'error': f'Ошибка чтения файла: {e}'},
            status=400
        )

    # ── Читаем заголовки (первая строка) ──────
    headers = {}
    for col_idx, cell in enumerate(
        next(ws.iter_rows(min_row=1, max_row=1)), start=0
    ):
        if cell.value:
            headers[str(cell.value).strip().lower()] = col_idx

    COLUMN_MAP = {
        'address':          ['адрес'],
        'cadastral_number': ['кадастровый номер', 'кадастровыйномер',
                             'кад. номер', 'кад.номер'],
        'customer':         ['заказчик'],
        'contract':         ['договор'],
        'newspaper':        ['газета'],
        'issue_date':       ['дата выпуска', 'датавыпуска'],
        'approval_date':    ['дата согласования', 'датасогласования'],
    }

    def find_col(key):
        for variant in COLUMN_MAP[key]:
            if variant in headers:
                return headers[variant]
        return None

    col = {key: find_col(key) for key in COLUMN_MAP}

    if col['address'] is None and col['cadastral_number'] is None:
        return JsonResponse(
            {'success': False,
             'error': 'Не найдены колонки «Адрес» или «Кадастровый номер». '
                      'Проверьте заголовки первой строки.'},
            status=400
        )

    # ── Читаем строки с "протяжкой" пустых значений ──────
    # Газета, дата выпуска, дата согласования — берём из предыдущей строки
    # если в текущей пусто (fill-down логика)
    rows_data = []

    prev_newspaper     = ''
    prev_issue_date    = None
    prev_approval_date = None

    for row in ws.iter_rows(min_row=2, values_only=True):

        def get_str(key, _row=row):
            idx = col[key]
            if idx is None:
                return ''
            val = _row[idx]
            return str(val).strip() if val not in (None, '') else ''

        def get_date(key, _row=row):
            idx = col[key]
            if idx is None:
                return None
            return parse_date(_row[idx])

        address          = get_str('address')
        cadastral_number = get_str('cadastral_number')
        customer         = get_str('customer')
        contract         = get_str('contract')

        # Пропускаем полностью пустые строки
        if not any([address, cadastral_number, customer]):
            continue

        # Газета и даты: если есть в текущей строке — обновляем prev,
        # если нет — используем предыдущее значение
        newspaper     = get_str('newspaper')
        issue_date    = get_date('issue_date')
        approval_date = get_date('approval_date')

        if newspaper:
            prev_newspaper = newspaper
        else:
            newspaper = prev_newspaper

        if issue_date:
            prev_issue_date = issue_date
        else:
            issue_date = prev_issue_date

        if approval_date:
            prev_approval_date = approval_date
        else:
            approval_date = prev_approval_date

        rows_data.append({
            'address':          address,
            'cadastral_number': cadastral_number,
            'customer':         customer,
            'contract':         contract,
            'newspaper':        newspaper,
            'issue_date':       issue_date,
            'approval_date':    approval_date,
        })

    if not rows_data:
        return JsonResponse(
            {'success': False, 'error': 'Файл не содержит данных'},
            status=400
        )

    # ── Группируем в извещения ────────────────
    # Новое извещение = смена газеты ИЛИ даты выпуска ИЛИ даты согласования
    groups = OrderedDict()
    for r in rows_data:
        key = (
            r['newspaper'] or '',
            str(r['issue_date']    or ''),
            str(r['approval_date'] or ''),
        )
        if key not in groups:
            groups[key] = {
                'newspaper':     r['newspaper'],
                'issue_date':    r['issue_date'],
                'approval_date': r['approval_date'],
                'items':         [],
            }
        groups[key]['items'].append(r)

    # ── Сохраняем в БД ────────────────────────
    created_notices = 0
    created_items   = 0

    for group in groups.values():
        notice = Notice.objects.create(
            newspaper=     group['newspaper']     or '',
            issue_date=    group['issue_date']    or None,
            approval_date= group['approval_date'] or None,
        )
        created_notices += 1

        for order, item in enumerate(group['items']):
            NoticeItem.objects.create(
                notice=           notice,
                address=          item['address'],
                cadastral_number= item['cadastral_number'],
                customer=         item['customer'],
                contract=         item['contract'],
                order=            order,
            )
            created_items += 1

    return JsonResponse({
        'success':         True,
        'created_notices': created_notices,
        'created_items':   created_items,
        'message':         (
            f'Импортировано: {created_notices} извещений, '
            f'{created_items} объектов'
        ),
    })


# ──────────────────────────────────────────────
# Шаблон Excel для скачивания
# ──────────────────────────────────────────────
def notice_import_template(request):
    """Скачать пустой шаблон Excel."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Извещения'

    headers = [
        'Адрес', 'Кадастровый номер', 'Заказчик',
        'Договор', 'Газета', 'Дата выпуска', 'Дата согласования',
    ]
    ws.append(headers)

    # Пример строки
    ws.append([
        'РФ, ЛНР, г. Луганск, ул. Примерная, 1',
        '95:19:0102065:1830',
        'физ. лицо',
        '01/КЗР-123 от 01.01.2025',
        'Республика',
        '04.09.2025',
        '06.10.2025',
    ])

    # ── Стиль заголовков ──────────────────────
    header_font = Font(bold=True, color='FFFFFF')
    header_fill = PatternFill('solid', fgColor='4472C4')

    for cell in ws[1]:          # ws[[1]](#annotation-142646-0) — первая строка, правильный синтаксис
        cell.font = header_font
        cell.fill = header_fill

    # ── Ширина колонок ────────────────────────
    col_widths = [50, 25, 20, 30, 15, 15, 20]
    for i, width in enumerate(col_widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = width

    response = HttpResponse(
        content_type=(
            'application/vnd.openxmlformats-officedocument'
            '.spreadsheetml.sheet'
        )
    )
    response['Content-Disposition'] = (
        'attachment; filename="notice_template.xlsx"'
    )
    wb.save(response)
    return response